from docx.table import _Cell

def reemplazar_placeholders_docx(doc_obj, reemplazos: dict):
    """
    Reemplaza todos los placeholders en un documento Word.
    Esta lógica está basada en la versión original y funcional del main.py.
    Maneja correctamente runs fragmentados y preserva el formato.
    """

    def reemplazar_en_parrafo(parrafo):
        texto_completo = ''.join(run.text for run in parrafo.runs)
        
        # Si no hay nada que reemplazar, salimos para no alterar el formato
        if not any(key in texto_completo for key in reemplazos.keys()):
            return

        # Realizamos todos los reemplazos en la cadena de texto
        for key, value in reemplazos.items():
            texto_completo = texto_completo.replace(key, str(value or ''))

        # Preservamos el formato del primer 'run' del párrafo
        if len(parrafo.runs) > 0:
            primer_run = parrafo.runs[0]
            fuente_nombre = primer_run.font.name
            fuente_tamanio = primer_run.font.size
            es_negrita = primer_run.bold
            es_cursiva = primer_run.italic
            es_subrayado = primer_run.underline

            # Limpiamos todos los runs del párrafo
            for run in parrafo.runs:
                run.text = ''
            
            # Añadimos un nuevo run con el texto completo y aplicamos el formato guardado
            nuevo_run = parrafo.add_run(texto_completo)
            if fuente_nombre:
                nuevo_run.font.name = fuente_nombre
            if fuente_tamanio:
                nuevo_run.font.size = fuente_tamanio
            if es_negrita is not None:
                nuevo_run.bold = es_negrita
            if es_cursiva is not None:
                nuevo_run.italic = es_cursiva
            if es_subrayado is not None:
                nuevo_run.underline = es_subrayado

    def reemplazar_en_tabla(tabla):
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    reemplazar_en_parrafo(parrafo)
                # Manejo de tablas anidadas
                for tabla_anidada in celda.tables:
                    reemplazar_en_tabla(tabla_anidada)

    # --- INICIO DEL PROCESO ---

    # Reemplazar en párrafos del cuerpo
    for parrafo in doc_obj.paragraphs:
        reemplazar_en_parrafo(parrafo)

    # Reemplazar en tablas del cuerpo
    for tabla in doc_obj.tables:
        reemplazar_en_tabla(tabla)

    # Reemplazar en encabezados y pies de página
    for seccion in doc_obj.sections:
        for parrafo in seccion.header.paragraphs:
            reemplazar_en_parrafo(parrafo)
        for tabla in seccion.header.tables:
            reemplazar_en_tabla(tabla)
        for parrafo in seccion.footer.paragraphs:
            reemplazar_en_parrafo(parrafo)
        for tabla in seccion.footer.tables:
            reemplazar_en_tabla(tabla)

